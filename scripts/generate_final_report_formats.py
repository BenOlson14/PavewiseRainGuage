from pathlib import Path
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

SRC = Path('docs/final-report.md')
DOCX_OUT = Path('docs/final-report.docx')
PDF_OUT = Path('docs/final-report.pdf')


def iter_markdown_lines(text: str):
    in_code = False
    for raw in text.splitlines():
        line = raw.rstrip('\n')
        if line.strip().startswith('```'):
            in_code = not in_code
            yield ('code_toggle', line)
            continue
        if in_code:
            yield ('code', line)
            continue
        if line.startswith('# '):
            yield ('h1', line[2:].strip())
        elif line.startswith('## '):
            yield ('h2', line[3:].strip())
        elif line.startswith('### '):
            yield ('h3', line[4:].strip())
        elif line.startswith('- '):
            yield ('bullet', line[2:].strip())
        elif line[:3].isdigit() and line[1:3] == '. ':
            yield ('number', line[3:].strip())
        elif line[:2].isdigit() and line[2:4] == ') ':
            yield ('number', line[4:].strip())
        elif line.strip() == '---':
            yield ('rule', '')
        else:
            yield ('text', line)


def to_docx(text: str):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    in_code = False
    for kind, payload in iter_markdown_lines(text):
        if kind == 'code_toggle':
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph(payload)
            p.style = doc.styles['No Spacing']
            if p.runs:
                p.runs[0].font.name = 'Courier New'
                p.runs[0].font.size = Pt(9)
            continue
        if kind == 'h1':
            doc.add_heading(payload, level=1)
        elif kind == 'h2':
            doc.add_heading(payload, level=2)
        elif kind == 'h3':
            doc.add_heading(payload, level=3)
        elif kind == 'bullet':
            doc.add_paragraph(payload, style='List Bullet')
        elif kind == 'number':
            doc.add_paragraph(payload, style='List Number')
        elif kind == 'rule':
            doc.add_paragraph('')
        else:
            doc.add_paragraph(payload)

    doc.save(DOCX_OUT)


def to_pdf(text: str):
    c = canvas.Canvas(str(PDF_OUT), pagesize=LETTER)
    width, height = LETTER
    margin = 54
    y = height - margin

    try:
        pdfmetrics.registerFont(TTFont('DejaVu', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
        body_font = 'DejaVu'
    except Exception:
        body_font = 'Helvetica'

    def new_page():
        nonlocal y
        c.showPage()
        y = height - margin

    in_code = False
    for kind, payload in iter_markdown_lines(text):
        if kind == 'code_toggle':
            in_code = not in_code
            continue

        if y < margin + 24:
            new_page()

        if in_code:
            c.setFont('Courier', 9)
            c.drawString(margin, y, payload[:120])
            y -= 12
            continue

        if kind == 'h1':
            c.setFont(body_font, 18)
            c.drawString(margin, y, payload)
            y -= 24
        elif kind == 'h2':
            c.setFont(body_font, 14)
            c.drawString(margin, y, payload)
            y -= 20
        elif kind == 'h3':
            c.setFont(body_font, 12)
            c.drawString(margin, y, payload)
            y -= 16
        elif kind == 'bullet':
            c.setFont(body_font, 10)
            c.drawString(margin + 8, y, f"• {payload}"[:130])
            y -= 14
        elif kind == 'number':
            c.setFont(body_font, 10)
            c.drawString(margin + 8, y, payload[:130])
            y -= 14
        elif kind == 'rule':
            c.line(margin, y, width - margin, y)
            y -= 12
        else:
            if not payload:
                y -= 10
                continue
            c.setFont(body_font, 10)
            # very simple wrapping
            words = payload.split()
            line = ''
            for w in words:
                test = f"{line} {w}".strip()
                if len(test) > 115:
                    c.drawString(margin, y, line)
                    y -= 13
                    if y < margin + 24:
                        new_page()
                    line = w
                else:
                    line = test
            if line:
                c.drawString(margin, y, line)
                y -= 13

    c.save()


def main():
    text = SRC.read_text(encoding='utf-8')
    to_docx(text)
    to_pdf(text)
    print(f'Wrote {DOCX_OUT} and {PDF_OUT}')


if __name__ == '__main__':
    main()
